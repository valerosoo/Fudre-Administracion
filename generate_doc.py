from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1C, 0x0A, 0x0E)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x7B, 0x2D, 0x3E)
    hs.font.name = 'Calibri'

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FUDRE ADMINISTRACION')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x7B, 0x2D, 0x3E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Documentacion Tecnica del Proyecto')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Sistema de administracion para club de vinos\nIntegracion con Tiendanube | Importacion inteligente con IA')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Junio 2026')
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INDICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Indice', level=1)
toc_items = [
    '1. Descripcion General',
    '2. Arquitectura del Sistema',
    '3. Tecnologias Utilizadas',
    '4. Estructura del Proyecto',
    '5. Backend - Spring Boot',
    '   5.1. Entidades y Base de Datos',
    '   5.2. Endpoints REST (API)',
    '   5.3. Servicios (Logica de Negocio)',
    '   5.4. Integracion con Tiendanube',
    '   5.5. Webhooks',
    '6. Frontend - React',
    '   6.1. Paginas',
    '   6.2. Servicios y Tipos',
    '7. Servicio de Importacion (IA)',
    '8. Base de Datos',
    '9. Configuracion y Variables de Entorno',
    '10. Como Levantar el Proyecto',
    '11. Flujos Principales',
    '12. Deployment (Railway)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. DESCRIPCION GENERAL
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Descripcion General', level=1)
doc.add_paragraph(
    'Fudre Administracion es un sistema de gestion interno para un club de vinos. '
    'Permite administrar el inventario de vinos, los miembros del club, las membresias '
    'y los envios mensuales. El sistema se integra automaticamente con la tienda online '
    'en Tiendanube (fudre.com.ar) para sincronizar productos, stock e imagenes.'
)
doc.add_paragraph(
    'Ademas, cuenta con un servicio de importacion inteligente que utiliza IA (Claude de Anthropic) '
    'para extraer datos de archivos PDF, Excel y CSV, permitiendo cargar informacion masiva '
    'sin necesidad de ingreso manual.'
)

doc.add_heading('Funcionalidades principales', level=2)
features = [
    'Gestion de vinos: alta, baja, modificacion, control de stock (gondola y cuartito), subida de fotos',
    'Gestion de miembros: datos personales, perfil de gustos, notas internas',
    'Gestion de membresias: planes BROTE, BROTE_PLUS, ENVERO, ENVERO_PLUS con estados activo/pausado/cancelado',
    'Gestion de envios: creacion con items, descuento automatico de stock, restauracion al eliminar',
    'Sincronizacion con Tiendanube: crear/editar/eliminar productos, subir imagenes, actualizar stock',
    'Webhooks de Tiendanube: recibir ordenes pagadas (descuenta stock), ordenes canceladas (restaura stock), clientes nuevos (crea miembro)',
    'Importacion inteligente: cargar datos desde PDF/Excel/CSV usando Claude AI con vista previa antes de confirmar',
    'Categoria automatica: los vinos se clasifican como BROTE (precio < $22.500) o ENVERO (>= $22.500)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Arquitectura del Sistema', level=1)
doc.add_paragraph(
    'El sistema esta compuesto por tres servicios independientes que se comunican entre si via HTTP:'
)

add_table(
    ['Servicio', 'Tecnologia', 'Puerto', 'Funcion'],
    [
        ['Backend', 'Spring Boot 4 (Java 21)', '8080', 'API REST, logica de negocio, integracion Tiendanube'],
        ['Frontend', 'React 19 + Vite + TypeScript', '5173', 'Panel de administracion (SPA)'],
        ['Import Service', 'Flask (Python)', '8081', 'Extraccion de datos con IA desde archivos'],
        ['Base de Datos', 'MySQL 8', '3306', 'Persistencia de datos'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'El frontend se comunica con el backend a traves de un proxy de Vite (/api -> localhost:8080). '
    'El servicio de importacion es llamado directamente desde el frontend y luego confirma los datos '
    'contra el backend.'
)

doc.add_paragraph()
doc.add_paragraph('Diagrama simplificado:', style='Intense Quote')
doc.add_paragraph(
    '  [Frontend :5173]  --->  [Backend :8080]  --->  [MySQL :3306]\n'
    '         |                      |      ^\n'
    '         v                      v      |\n'
    '  [Import Service :8081]   [Tiendanube API]\n'
    '         |                      ^\n'
    '         v                      |\n'
    '    [Claude AI API]       [Webhooks]'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. TECNOLOGIAS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Tecnologias Utilizadas', level=1)

doc.add_heading('Backend', level=2)
add_table(
    ['Tecnologia', 'Version', 'Uso'],
    [
        ['Java', '21', 'Lenguaje principal'],
        ['Spring Boot', '4.0.6', 'Framework web y DI'],
        ['Spring Data JPA', '-', 'ORM y repositorios'],
        ['Flyway', '-', 'Migraciones de base de datos'],
        ['MySQL Connector', '-', 'Driver de base de datos'],
        ['Gradle', '-', 'Build system'],
    ]
)

doc.add_paragraph()
doc.add_heading('Frontend', level=2)
add_table(
    ['Tecnologia', 'Version', 'Uso'],
    [
        ['React', '19.2.6', 'UI framework'],
        ['TypeScript', '6.0.2', 'Tipado estatico'],
        ['Vite', '8.0.12', 'Build tool y dev server'],
        ['Tailwind CSS', '4.3.0', 'Estilos utilitarios'],
        ['React Router', '7.15.1', 'Navegacion SPA'],
        ['React Hook Form', '7.76.0', 'Formularios'],
        ['Radix UI / Shadcn', '-', 'Componentes UI accesibles'],
        ['Sonner', '2.0.7', 'Notificaciones toast'],
        ['Lucide React', '-', 'Iconos'],
    ]
)

doc.add_paragraph()
doc.add_heading('Servicio de Importacion', level=2)
add_table(
    ['Tecnologia', 'Uso'],
    [
        ['Python 3 + Flask', 'Servidor HTTP'],
        ['Anthropic SDK', 'Llamadas a Claude AI'],
        ['pdfplumber', 'Extraccion de texto de PDFs'],
        ['pandas + openpyxl', 'Lectura de Excel/CSV'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. ESTRUCTURA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Estructura del Proyecto', level=1)
doc.add_paragraph(
    'El proyecto sigue una estructura de monorepo con tres carpetas principales:'
)

structure = """Fudre-Administracion/
|-- backend/
|   |-- build.gradle
|   |-- src/main/java/fudre/app/
|   |   |-- AppApplication.java
|   |   |-- config/
|   |   |   |-- CorsConfig.java
|   |   |   |-- FlywayConfig.java
|   |   |   |-- TiendanubeConfig.java
|   |   |   |-- WebMvcConfig.java
|   |   |-- controller/
|   |   |   |-- MemberController.java
|   |   |   |-- MembershipController.java
|   |   |   |-- ShipmentController.java
|   |   |   |-- TiendanubeWebhookController.java
|   |   |   |-- WineController.java
|   |   |-- dto/
|   |   |   |-- MemberDto, MembershipDto, ShipmentDto, ShipmentItemDto, WineDto
|   |   |-- entity/
|   |   |   |-- Member, MemberGrapeRating, Membership, Shipment, ShipmentItem
|   |   |   |-- Wine, WineCategory, WinePool, WineStyle, Plan, MembershipStatus
|   |   |-- repository/
|   |   |   |-- MemberRepository, MembershipRepository, ShipmentRepository
|   |   |   |-- WineRepository, WinePoolRepository
|   |   |-- service/
|   |       |-- MemberService, MembershipService, ShipmentService
|   |       |-- StockSyncService, TiendanubeService, WineService
|   |-- src/main/resources/
|       |-- application.properties
|       |-- application-railway.properties
|       |-- db/migration/
|           |-- V1__init.sql ... V5__add_wine_image_url.sql
|-- frontend/
|   |-- package.json, vite.config.ts, tsconfig.json
|   |-- src/
|       |-- App.tsx, main.tsx, index.css
|       |-- components/
|       |   |-- Layout.tsx, ImportButton.tsx
|       |   |-- ui/ (button, input, dialog, table, badge, etc.)
|       |-- pages/
|       |   |-- WinesPage.tsx, MembersPage.tsx
|       |   |-- MembershipsPage.tsx, ShipmentsPage.tsx
|       |-- services/
|       |   |-- api.ts, wines.ts, members.ts, memberships.ts, shipments.ts
|       |-- types/index.ts
|-- import_service/
|   |-- main.py, config.py, ai_client.py, prompts.py, extractors.py
|-- load_data.py
|-- README.md"""

p = doc.add_paragraph(structure)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. BACKEND
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Backend - Spring Boot', level=1)
doc.add_paragraph(
    'El backend es una aplicacion Spring Boot 4 que expone una API REST bajo /api. '
    'Usa JPA/Hibernate para la persistencia y Flyway para migraciones incrementales de la base de datos.'
)

# 5.1 Entidades
doc.add_heading('5.1. Entidades y Base de Datos', level=2)

doc.add_heading('Wine (Vino)', level=3)
add_table(
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['id', 'Long (PK, auto)', 'Identificador unico'],
        ['name', 'String (not null)', 'Nombre del vino (ej: "Malbec Reserva")'],
        ['grape', 'String', 'Cepa (ej: "Malbec")'],
        ['vintageYear', 'Integer', 'Anio de cosecha'],
        ['stockGondola', 'Integer (default 0)', 'Unidades en gondola'],
        ['stockCuartito', 'Integer (default 0)', 'Unidades en cuartito (deposito)'],
        ['referencePrice', 'BigDecimal', 'Precio de referencia en ARS'],
        ['category', 'Enum (generado)', 'BROTE si precio < 22500, ENVERO si >= 22500'],
        ['isClubEligible', 'Boolean (default false)', 'Si esta habilitado para el club'],
        ['tiendanubeProductId', 'String', 'ID del producto en Tiendanube'],
        ['tiendanubeVariantId', 'String', 'ID de la variante en Tiendanube'],
        ['imageUrl', 'String', 'URL de la imagen del vino'],
        ['uploadStatus', 'String', 'PENDING | UPLOADED | OUT_OF_STOCK'],
        ['createdAt', 'LocalDateTime', 'Fecha de creacion'],
    ]
)

doc.add_paragraph()
doc.add_heading('Member (Miembro)', level=3)
add_table(
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['id', 'Long (PK, auto)', 'Identificador unico'],
        ['name', 'String (not null)', 'Nombre completo'],
        ['email', 'String (unique, not null)', 'Email del miembro'],
        ['phone', 'String', 'Telefono'],
        ['deliveryAddress', 'String', 'Direccion de entrega'],
        ['wineStyle', 'Enum', 'JOVENES | MAS_CUERPO'],
        ['wineTypes', 'String', 'Tipos de vino preferidos'],
        ['openToNew', 'Boolean', 'Abierto a probar nuevos'],
        ['occasions', 'String', 'Ocasiones de consumo'],
        ['createdAt', 'LocalDateTime', 'Fecha de creacion'],
    ]
)

doc.add_paragraph()
doc.add_heading('Membership (Membresia)', level=3)
add_table(
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['id', 'Long (PK, auto)', 'Identificador unico'],
        ['member', 'Member (FK)', 'Miembro asociado'],
        ['plan', 'Enum', 'BROTE | BROTE_PLUS | ENVERO | ENVERO_PLUS'],
        ['status', 'Enum (default ACTIVE)', 'ACTIVE | PAUSED | CANCELLED'],
        ['startDate', 'LocalDate', 'Fecha de inicio'],
        ['endDate', 'LocalDate', 'Fecha de fin (opcional)'],
    ]
)

doc.add_paragraph()
doc.add_heading('Shipment (Envio)', level=3)
add_table(
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['id', 'Long (PK, auto)', 'Identificador unico'],
        ['member', 'Member (FK)', 'Miembro destinatario'],
        ['membership', 'Membership (FK)', 'Membresia asociada'],
        ['shippedAt', 'LocalDate', 'Fecha de envio'],
        ['shippingCost', 'BigDecimal', 'Costo de envio'],
        ['notes', 'String', 'Notas internas'],
        ['tiendanubeOrderId', 'String (unique)', 'ID de orden Tiendanube (idempotencia)'],
        ['items', 'List<ShipmentItem>', 'Vinos incluidos en el envio'],
    ]
)

doc.add_paragraph()
doc.add_heading('ShipmentItem (Item de Envio)', level=3)
add_table(
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['id', 'Long (PK, auto)', 'Identificador unico'],
        ['shipment', 'Shipment (FK)', 'Envio al que pertenece'],
        ['wine', 'Wine (FK)', 'Vino enviado'],
        ['quantity', 'Integer (default 1)', 'Cantidad'],
        ['unitPrice', 'BigDecimal', 'Precio unitario al momento del envio'],
    ]
)

doc.add_page_break()

# 5.2 Endpoints
doc.add_heading('5.2. Endpoints REST (API)', level=2)

doc.add_heading('Vinos - /api/wines', level=3)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['GET', '/api/wines', 'Listar todos los vinos'],
        ['GET', '/api/wines/{id}', 'Obtener vino por ID'],
        ['POST', '/api/wines', 'Crear vino (sincroniza con Tiendanube)'],
        ['PUT', '/api/wines/{id}', 'Actualizar vino (sincroniza con Tiendanube)'],
        ['DELETE', '/api/wines/{id}', 'Eliminar vino (elimina de Tiendanube)'],
        ['POST', '/api/wines/{id}/image', 'Subir imagen (multipart/form-data, campo "file")'],
    ]
)

doc.add_paragraph()
doc.add_heading('Miembros - /api/members', level=3)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['GET', '/api/members', 'Listar todos los miembros'],
        ['GET', '/api/members/{id}', 'Obtener miembro por ID'],
        ['POST', '/api/members', 'Crear miembro'],
        ['PUT', '/api/members/{id}', 'Actualizar miembro'],
        ['DELETE', '/api/members/{id}', 'Eliminar miembro'],
    ]
)

doc.add_paragraph()
doc.add_heading('Membresias - /api/memberships', level=3)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['GET', '/api/memberships', 'Listar todas las membresias'],
        ['GET', '/api/memberships/{id}', 'Obtener membresia por ID'],
        ['GET', '/api/memberships/member/{memberId}', 'Membresias de un miembro'],
        ['POST', '/api/memberships', 'Crear membresia'],
        ['PUT', '/api/memberships/{id}', 'Actualizar membresia'],
        ['DELETE', '/api/memberships/{id}', 'Eliminar membresia'],
    ]
)

doc.add_paragraph()
doc.add_heading('Envios - /api/shipments', level=3)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['GET', '/api/shipments', 'Listar todos los envios'],
        ['GET', '/api/shipments/{id}', 'Obtener envio por ID'],
        ['GET', '/api/shipments/member/{memberId}', 'Envios de un miembro'],
        ['POST', '/api/shipments', 'Crear envio (descuenta stock automaticamente)'],
        ['DELETE', '/api/shipments/{id}', 'Eliminar envio (restaura stock)'],
    ]
)

doc.add_paragraph()
doc.add_heading('Webhooks - /api/webhooks/tiendanube', level=3)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['POST', '/api/webhooks/tiendanube', 'Endpoint unificado para todos los eventos'],
        ['POST', '/api/webhooks/tiendanube/orders', 'Endpoint legacy para ordenes'],
    ]
)

doc.add_page_break()

# 5.3 Servicios
doc.add_heading('5.3. Servicios (Logica de Negocio)', level=2)

doc.add_heading('WineService', level=3)
doc.add_paragraph(
    'Gestiona el CRUD de vinos y la sincronizacion con Tiendanube.'
)
services_wine = [
    'create(): Guarda el vino en DB, luego intenta crear el producto en Tiendanube. Si tiene exito, guarda el productId y variantId y marca como UPLOADED. Si falla, queda como PENDING.',
    'update(): Actualiza el vino en DB, luego intenta actualizar nombre/precio/stock en Tiendanube.',
    'delete(): Intenta eliminar el producto de Tiendanube, luego elimina de DB (siempre).',
    'uploadImage(): Guarda la imagen en disco (uploads/wines/), actualiza imageUrl en DB, y envia la imagen a Tiendanube como src URL.',
]
for s in services_wine:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('ShipmentService', level=3)
doc.add_paragraph(
    'Gestiona la creacion y eliminacion de envios con descuento/restauracion automatica de stock.'
)
services_ship = [
    'create(): Crea el envio con sus items, descuenta stock por cada item (primero de gondola, luego de cuartito). Sincroniza stock actualizado a Tiendanube.',
    'delete(): Restaura el stock de cada item y elimina el envio.',
]
for s in services_ship:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('StockSyncService', level=3)
doc.add_paragraph(
    'Maneja la sincronizacion de stock entre ordenes de Tiendanube y el sistema local.'
)
services_stock = [
    'discountStockFromOrder(): Busca los vinos por tiendanubeProductId y descuenta las cantidades compradas.',
    'restoreStockFromOrder(): Restaura el stock cuando una orden es cancelada.',
    'createShipmentFromOrder(): Crea un envio automatico a partir de una orden pagada (con idempotencia por tiendanubeOrderId).',
    'syncStockToTiendanube(): Envia el stock total actualizado a Tiendanube via API.',
]
for s in services_stock:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('MemberService', level=3)
doc.add_paragraph(
    'CRUD de miembros. Incluye createOrSkip() para webhooks: si el email ya existe, no crea duplicado.'
)

doc.add_page_break()

# 5.4 Tiendanube
doc.add_heading('5.4. Integracion con Tiendanube', level=2)
doc.add_paragraph(
    'La integracion usa la API REST de Tiendanube (https://api.tiendanube.com/v1/{storeId}). '
    'Se configura con store-id y access-token (obtenido via OAuth). El token es permanente una vez obtenido.'
)

doc.add_heading('TiendanubeService - Metodos', level=3)
add_table(
    ['Metodo', 'Endpoint Tiendanube', 'Descripcion'],
    [
        ['createProduct()', 'POST /products', 'Crea producto con nombre, precio y stock. Retorna [productId, variantId]'],
        ['updateProduct()', 'PUT /products/{id} + PUT /variants/{id}', 'Actualiza nombre, precio y stock'],
        ['deleteProduct()', 'DELETE /products/{id}', 'Elimina producto'],
        ['updateStock()', 'PUT /products/{id}/variants/{id}', 'Actualiza solo el stock'],
        ['uploadProductImage()', 'POST /products/{id}/images', 'Sube imagen via URL src'],
        ['getOrder()', 'GET /orders/{id}', 'Obtiene detalle de una orden'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Nota sobre imagenes: Tiendanube requiere una URL publica para descargar la imagen. '
    'En desarrollo local (localhost) las imagenes no se veran en la tienda. '
    'En produccion, configurar fudre.server.base-url con la URL publica del servidor.'
)

# 5.5 Webhooks
doc.add_heading('5.5. Webhooks de Tiendanube', level=2)
doc.add_paragraph(
    'El controlador TiendanubeWebhookController recibe eventos POST de Tiendanube:'
)
add_table(
    ['Evento', 'Accion del Sistema'],
    [
        ['order/paid', 'Obtiene la orden, descuenta stock de los vinos comprados, crea un envio automatico'],
        ['order/cancelled', 'Restaura el stock descontado, elimina el envio asociado'],
        ['customer/created', 'Crea un miembro nuevo si el email no existe'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Los webhooks se configuran en Tiendanube Partners apuntando a: '
    'https://tu-servidor/api/webhooks/tiendanube'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. FRONTEND
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Frontend - React', level=1)
doc.add_paragraph(
    'El frontend es una SPA (Single Page Application) con React 19, TypeScript y Tailwind CSS. '
    'Usa Vite como build tool y dev server con proxy hacia el backend.'
)

doc.add_heading('Layout', level=2)
doc.add_paragraph(
    'La aplicacion tiene un header fijo superior (fondo oscuro) con saludo y fecha/hora, '
    'un sidebar lateral izquierdo (color bordo/marron #7F654E) con iconos de navegacion, '
    'y un area principal con fondo blanco. Los colores de marca son: bordo profundo (#7B2D3E), '
    'dorado (#C9A84C) y crema (#FAF8F4).'
)

# 6.1 Paginas
doc.add_heading('6.1. Paginas', level=2)

doc.add_heading('WinesPage (Vinos)', level=3)
doc.add_paragraph(
    'Tabla con columnas: Foto, Nombre, Uva, Anio, Stock Gondola, Stock Cuartito, Total, '
    'Precio Ref., Categoria (badge), Club (badge), Estado Tiendanube, Acciones (editar/eliminar).'
)
doc.add_paragraph('Filtros disponibles:', style='List Bullet')
wine_filters = [
    'Busqueda por nombre',
    'Ordenar por nombre (A-Z / Z-A)',
    'Filtrar por uva (multi-seleccion)',
    'Filtrar por stock (todo / con stock / sin stock)',
    'Filtrar por categoria (BROTE / ENVERO)',
    'Filtrar por estado de subida (Pendiente / Subido / Sin stock)',
]
for f in wine_filters:
    doc.add_paragraph(f, style='List Bullet 2')
doc.add_paragraph(
    'El formulario de creacion/edicion incluye: nombre, uva, anio, stock gondola/cuartito, '
    'precio referencia, checkbox "apto para club", y selector de imagen con vista previa.'
)

doc.add_heading('MembersPage (Miembros)', level=3)
doc.add_paragraph(
    'Tabla con: Nombre, Email, Telefono, Direccion, Perfil de Gustos, Acciones. '
    'Filtro de busqueda por nombre. Formulario con campos de texto y areas para notas.'
)

doc.add_heading('MembershipsPage (Membresias)', level=3)
doc.add_paragraph(
    'Tabla con: Nombre del Miembro, Plan (badge), Fecha Inicio, Estado (badge), Acciones. '
    'Filtros por nombre de miembro y por plan. Formulario con seleccion de miembro, plan y fecha.'
)

doc.add_heading('ShipmentsPage (Envios)', level=3)
doc.add_paragraph(
    'Tabla con: Miembro, Email, Fecha, Costo Envio, Cantidad Items, Notas, Eliminar. '
    'Filtros por nombre y ordenar por fecha. Formulario de creacion con seleccion de miembro, '
    'membresia, fecha, items dinamicos (vino + cantidad + precio), costo de envio y notas. '
    'Al eliminar se restaura automaticamente el stock.'
)

# 6.2 Servicios
doc.add_heading('6.2. Servicios y Tipos', level=2)
doc.add_paragraph(
    'Todos los servicios frontend estan en src/services/ y usan un wrapper HTTP base (api.ts) '
    'que maneja errores y serializa JSON automaticamente. El base path es /api (proxied a :8080).'
)
add_table(
    ['Servicio', 'Metodos'],
    [
        ['wines.ts', 'getAll(), create(), update(), delete(), uploadImage()'],
        ['members.ts', 'getAll(), create(), update(), delete()'],
        ['memberships.ts', 'getAll(), create(), update(), delete()'],
        ['shipments.ts', 'getAll(), create(), delete()'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Los tipos TypeScript estan definidos en src/types/index.ts e incluyen: '
    'Wine, Member, Membership, Shipment, ShipmentItem, y los enums Plan, WineCategory, UploadStatus.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. IMPORT SERVICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. Servicio de Importacion (IA)', level=1)
doc.add_paragraph(
    'Microservicio Python (Flask) que permite importar datos masivamente desde archivos. '
    'Usa Claude AI (modelo claude-haiku-4-5) para extraer informacion estructurada de documentos '
    'no estandarizados.'
)

doc.add_heading('Flujo de importacion', level=2)
steps = [
    '1. El usuario sube un archivo (PDF, Excel, CSV) desde el boton "Importar" en el frontend.',
    '2. El archivo se envia al import_service (POST /import/{entity}).',
    '3. El servicio extrae el texto del archivo (pdfplumber para PDFs, pandas para Excel/CSV).',
    '4. Se envia el texto a Claude AI con un prompt especifico para la entidad.',
    '5. Claude retorna un JSON estructurado con los datos extraidos.',
    '6. El servicio devuelve una vista previa al frontend.',
    '7. El usuario revisa los datos y confirma la importacion.',
    '8. El frontend envia los datos confirmados a POST /import/{entity}/confirm.',
    '9. El import_service los inserta uno a uno via la API del backend.',
]
for s in steps:
    doc.add_paragraph(s)

doc.add_heading('Endpoints del Import Service', level=2)
add_table(
    ['Metodo', 'Ruta', 'Descripcion'],
    [
        ['GET', '/health', 'Health check'],
        ['POST', '/import/{entity}', 'Parsear archivo y devolver preview (sin insertar)'],
        ['POST', '/import/{entity}/confirm', 'Confirmar e insertar los datos en el backend'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Entidades soportadas: wines, members, memberships, shipments.')

doc.add_heading('Configuracion', level=2)
add_table(
    ['Variable de Entorno', 'Default', 'Descripcion'],
    [
        ['ANTHROPIC_API_KEY', '(requerida)', 'API key de Anthropic para Claude'],
        ['CLAUDE_MODEL', 'claude-haiku-4-5-20251001', 'Modelo de Claude a usar'],
        ['BACKEND_URL', 'http://localhost:8080/api', 'URL de la API del backend'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. BASE DE DATOS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Base de Datos', level=1)
doc.add_paragraph(
    'MySQL 8. La base de datos se llama fudre_admin. Las tablas se crean y actualizan '
    'automaticamente con Flyway al iniciar el backend.'
)

doc.add_heading('Migraciones', level=2)
add_table(
    ['Migracion', 'Descripcion'],
    [
        ['V1__init.sql', 'Crea todas las tablas base: members, member_grape_ratings, memberships, wines, wine_pool, shipments, shipment_items. La columna category de wines es generada automaticamente por precio.'],
        ['V2__fix_envero_typo.sql', 'Corrige typo EMVERO -> ENVERO en el enum de categoria'],
        ['V3__add_tiendanube_order_id.sql', 'Agrega tiendanube_order_id a shipments con indice unico (idempotencia)'],
        ['V4__add_tiendanube_variant_id.sql', 'Agrega tiendanube_variant_id a wines'],
        ['V5__add_wine_image_url.sql', 'Agrega image_url a wines'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Nota: La columna category de la tabla wines es una columna generada (GENERATED ALWAYS AS). '
    'Si reference_price < 22500 es BROTE, si >= 22500 es ENVERO. No se puede insertar ni '
    'actualizar manualmente.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. CONFIGURACION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Configuracion y Variables de Entorno', level=1)

doc.add_heading('application.properties (desarrollo local)', level=2)
add_table(
    ['Propiedad', 'Valor por defecto', 'Descripcion'],
    [
        ['spring.datasource.url', 'jdbc:mysql://localhost:3306/fudre_admin', 'URL de la base de datos'],
        ['spring.datasource.username', 'root', 'Usuario MySQL'],
        ['spring.datasource.password', '(tu password)', 'Password MySQL'],
        ['tiendanube.store-id', '(tu store id)', 'ID de la tienda en Tiendanube'],
        ['tiendanube.access-token', '(tu token)', 'Token de acceso OAuth permanente'],
        ['fudre.uploads.dir', 'uploads', 'Directorio para imagenes subidas'],
        ['fudre.server.base-url', 'http://localhost:8080', 'URL publica del servidor (para imagenes en Tiendanube)'],
    ]
)

doc.add_paragraph()
doc.add_heading('application-railway.properties (produccion)', level=2)
doc.add_paragraph(
    'En produccion (Railway) todas las credenciales se leen de variables de entorno:'
)
add_table(
    ['Variable', 'Descripcion'],
    [
        ['DATABASE_URL', 'URL JDBC completa de la base MySQL'],
        ['DB_USERNAME', 'Usuario de la base de datos'],
        ['DB_PASSWORD', 'Password de la base de datos'],
        ['TIENDANUBE_STORE_ID', 'ID de la tienda'],
        ['TIENDANUBE_ACCESS_TOKEN', 'Token OAuth permanente'],
        ['RAILWAY_PUBLIC_DOMAIN', 'URL publica del servidor (con https://)'],
    ]
)

doc.add_paragraph()
doc.add_heading('Proxy de Vite (vite.config.ts)', level=2)
doc.add_paragraph(
    'En desarrollo, Vite proxea dos rutas al backend:\n'
    '  /api -> http://localhost:8080\n'
    '  /uploads -> http://localhost:8080\n'
    'Esto evita problemas de CORS y permite usar rutas relativas.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. COMO LEVANTAR
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Como Levantar el Proyecto', level=1)

doc.add_heading('Requisitos previos', level=2)
reqs = [
    'Java 21 (JDK)',
    'Node.js 18+ y npm',
    'Python 3.10+',
    'MySQL 8 corriendo en localhost:3306',
    'Crear la base de datos: CREATE DATABASE fudre_admin;',
]
for r in reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Paso 1: Backend', level=2)
doc.add_paragraph(
    'cd backend\n'
    './gradlew bootRun'
)
doc.add_paragraph('El backend inicia en http://localhost:8080. Flyway ejecuta las migraciones automaticamente.')

doc.add_heading('Paso 2: Frontend', level=2)
doc.add_paragraph(
    'cd frontend\n'
    'npm install\n'
    'npm run dev'
)
doc.add_paragraph('El frontend inicia en http://localhost:5173.')

doc.add_heading('Paso 3: Servicio de Importacion (opcional)', level=2)
doc.add_paragraph(
    'cd import_service\n'
    'pip install flask flask-cors pandas openpyxl pdfplumber requests anthropic\n'
    'set ANTHROPIC_API_KEY=tu-api-key\n'
    'python main.py'
)
doc.add_paragraph('El servicio de importacion inicia en http://localhost:8081.')

doc.add_heading('Verificacion', level=2)
doc.add_paragraph(
    'Abrir http://localhost:5173 en el navegador. Deberia mostrar el panel de administracion '
    'con la seccion de vinos. Si el backend esta corriendo, la tabla muestra los vinos de la base de datos.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. FLUJOS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Flujos Principales', level=1)

doc.add_heading('Crear un vino', level=2)
doc.add_paragraph(
    '1. El admin completa el formulario en el frontend (nombre, uva, precio, stock, imagen).\n'
    '2. Frontend envia POST /api/wines con los datos.\n'
    '3. Backend guarda el vino en DB con uploadStatus = PENDING.\n'
    '4. Backend intenta POST /products en Tiendanube.\n'
    '5. Si Tiendanube responde OK: guarda productId + variantId, marca UPLOADED.\n'
    '6. Si falla: queda como PENDING (el vino NO se pierde).\n'
    '7. Si el admin subio una imagen, el frontend envia POST /api/wines/{id}/image.\n'
    '8. Backend guarda la imagen en disco y envia la URL a Tiendanube.'
)

doc.add_heading('Crear un envio', level=2)
doc.add_paragraph(
    '1. El admin selecciona miembro, membresia, fecha y agrega vinos con cantidades.\n'
    '2. Frontend envia POST /api/shipments.\n'
    '3. Backend crea el envio y por cada item descuenta stock (primero gondola, luego cuartito).\n'
    '4. Si el vino tiene tiendanubeProductId, sincroniza el stock nuevo a Tiendanube.\n'
    '5. Si se elimina el envio, el stock se restaura automaticamente.'
)

doc.add_heading('Orden pagada en Tiendanube (webhook)', level=2)
doc.add_paragraph(
    '1. Un cliente compra en fudre.com.ar.\n'
    '2. Tiendanube envia POST /api/webhooks/tiendanube con event = "order/paid".\n'
    '3. Backend obtiene los detalles de la orden via GET /orders/{id}.\n'
    '4. Descuenta stock de los vinos comprados.\n'
    '5. Crea un envio automatico (si no existe ya por idempotencia).\n'
    '6. Si la orden se cancela despues, el webhook order/cancelled restaura el stock.'
)

doc.add_heading('Importar datos con IA', level=2)
doc.add_paragraph(
    '1. El admin hace clic en "Importar" y sube un archivo PDF/Excel/CSV.\n'
    '2. El archivo se envia al import_service.\n'
    '3. Claude AI extrae los datos y los retorna como JSON.\n'
    '4. El frontend muestra una tabla de preview.\n'
    '5. El admin revisa y confirma.\n'
    '6. Los datos se insertan uno a uno via la API del backend.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Deployment', level=1)

doc.add_heading('Railway', level=2)
doc.add_paragraph(
    'El backend esta preparado para deploy en Railway con el perfil "railway" '
    '(application-railway.properties). Todas las credenciales se leen de variables de entorno.'
)

doc.add_heading('Variables de entorno necesarias en Railway', level=3)
add_table(
    ['Variable', 'Ejemplo'],
    [
        ['SPRING_PROFILES_ACTIVE', 'railway'],
        ['DATABASE_URL', 'jdbc:mysql://host:3306/dbname'],
        ['DB_USERNAME', 'root'],
        ['DB_PASSWORD', '***'],
        ['TIENDANUBE_STORE_ID', '1551637'],
        ['TIENDANUBE_ACCESS_TOKEN', '***'],
        ['RAILWAY_PUBLIC_DOMAIN', 'https://tu-app.up.railway.app'],
    ]
)

doc.add_paragraph()
doc.add_heading('Servidor propio', level=2)
doc.add_paragraph(
    'Para deployar en un servidor propio (VPS, dedicado, etc.):'
)
deploy_steps = [
    '1. Instalar Java 21, MySQL 8, Node.js 18+, Python 3.10+ en el servidor.',
    '2. Crear la base de datos MySQL: CREATE DATABASE fudre_admin;',
    '3. Clonar el repositorio en el servidor.',
    '4. Backend: ejecutar ./gradlew bootJar para generar el JAR. Correr con: java -jar backend/build/libs/app.jar --spring.profiles.active=prod',
    '5. Frontend: ejecutar npm run build en /frontend. Servir la carpeta dist/ con nginx o similar.',
    '6. Configurar las variables de entorno o crear un application-prod.properties con las credenciales.',
    '7. Configurar fudre.server.base-url con la URL publica (ej: https://admin.fudre.com.ar).',
    '8. Configurar nginx como reverse proxy: puerto 80/443 -> localhost:8080 (backend) y servir /dist (frontend).',
    '9. Configurar los webhooks de Tiendanube con la URL publica del servidor.',
]
for s in deploy_steps:
    doc.add_paragraph(s)

# ═══════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════
output_path = r'c:\Users\Pasante\Desktop\Fudre\Fudre-Administracion\Fudre_Administracion_Documentacion.docx'
doc.save(output_path)
print(f'Documento generado en: {output_path}')
